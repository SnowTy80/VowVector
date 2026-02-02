"""Text extraction from supported file types.

Dispatches to type-specific handlers. Falls back to OCR for scanned
PDFs and image files via core/ocr.py.
"""

import csv
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path

from config import MAX_CONTENT_CHARS, OCR_DPI, SUPPORTED_EXTENSIONS

logger = logging.getLogger(__name__)


@dataclass
class ExtractedDocument:
    """Result of text extraction from a single file."""

    source_file: str
    raw_text: str = ""
    page_count: int = 1
    extraction_method: str = "native_text"
    warnings: list[str] = field(default_factory=list)


def extract_text(
    file_path: Path,
    use_ocr: bool = True,
    ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Main extraction dispatcher.

    Args:
        file_path: Path to the input file.
        use_ocr: Whether to attempt OCR on image-based content.
        ocr_engine: "tesseract" or "nanonets".
        force_ocr: If True, skip native text extraction and OCR every page
                   (for graphical PDFs like construction drawings).

    Returns:
        ExtractedDocument with extracted text and metadata.

    Raises:
        ValueError: If the file type is unsupported.
    """
    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    handler = _EXTRACTORS.get(ext)
    if handler is None:
        raise ValueError(f"No extraction handler for: {ext}")

    doc = handler(file_path, use_ocr=use_ocr, ocr_engine=ocr_engine, force_ocr=force_ocr)

    # Safety truncation for extremely large documents
    if len(doc.raw_text) > MAX_CONTENT_CHARS:
        doc.raw_text = doc.raw_text[:MAX_CONTENT_CHARS]
        doc.warnings.append(
            f"Text truncated to {MAX_CONTENT_CHARS:,} characters"
        )

    return doc


# ── PDF ──

def _extract_pdf(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Extract text from PDF using PyMuPDF.

    Tries native text extraction first. If the average characters per page
    is below 50 (likely a scanned document), falls back to OCR.
    If force_ocr is True, skips native text and OCRs every page (for
    graphical PDFs like construction drawings).
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    page_count = len(doc)

    extraction_method = "native_text"
    warnings = []
    needs_ocr = force_ocr

    if not force_ocr:
        pages_text = []
        for page in doc:
            pages_text.append(page.get_text())
        total_chars = sum(len(t) for t in pages_text)
        avg_chars = total_chars / max(page_count, 1)
        if avg_chars < 50 and use_ocr:
            needs_ocr = True
    else:
        pages_text = [""] * page_count

    # OCR path
    if needs_ocr and use_ocr:
        from core.ocr import ocr_page

        extraction_method = f"ocr_{ocr_engine}"
        ocr_pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=OCR_DPI)
            img_bytes = pix.tobytes("png")
            from PIL import Image

            img = Image.open(io.BytesIO(img_bytes))
            try:
                text = ocr_page(img, engine=ocr_engine)
                ocr_pages.append(text)
            except Exception as e:
                warnings.append(f"OCR failed on page {page.number + 1}: {e}")
                ocr_pages.append(pages_text[page.number])

        pages_text = ocr_pages

    doc.close()
    full_text = "\n\n".join(pages_text).strip()

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text=full_text,
        page_count=page_count,
        extraction_method=extraction_method,
        warnings=warnings,
    )


# ── DOCX ──

def _extract_docx(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Extract text from .docx files using python-docx.

    Handles paragraphs and tables.
    """
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            parts.append("\n".join(rows_text))

    warnings = []
    if file_path.suffix.lower() == ".doc":
        warnings.append(
            "Legacy .doc format — python-docx may not fully parse this file. "
            "Consider converting to .docx for best results."
        )

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text="\n\n".join(parts).strip(),
        page_count=1,  # python-docx doesn't expose page count
        extraction_method="native_text",
        warnings=warnings,
    )


# ── XLSX ──

def _extract_xlsx(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Extract text from Excel files using openpyxl.

    Iterates all sheets, converts rows to pipe-separated text.
    """
    from openpyxl import load_workbook

    wb = load_workbook(str(file_path), read_only=True, data_only=True)
    parts = []
    warnings = []

    if file_path.suffix.lower() == ".xls":
        warnings.append(
            "Legacy .xls format — openpyxl may not parse this file. "
            "Consider converting to .xlsx for best results."
        )

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = " | ".join(cells).strip()
            if line and line != "|".join([""] * len(cells)).strip():
                sheet_rows.append(line)
        if sheet_rows:
            parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(sheet_rows))

    wb.close()

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text="\n\n".join(parts).strip(),
        page_count=len(wb.sheetnames) if parts else 1,
        extraction_method="table_parse",
        warnings=warnings,
    )


# ── CSV ──

def _extract_csv(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Read CSV and join rows into text blocks."""
    rows_text = []
    warnings = []

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                line = " | ".join(row).strip()
                if line:
                    rows_text.append(line)
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            reader = csv.reader(f)
            for row in reader:
                line = " | ".join(row).strip()
                if line:
                    rows_text.append(line)
        warnings.append("File decoded with latin-1 fallback")

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text="\n".join(rows_text).strip(),
        page_count=1,
        extraction_method="table_parse",
        warnings=warnings,
    )


# ── TXT / MD ──

def _extract_txt(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """Read plain text with UTF-8/Latin-1 fallback."""
    warnings = []

    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
        warnings.append("File decoded with latin-1 fallback")

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text=text.strip(),
        page_count=1,
        extraction_method="native_text",
        warnings=warnings,
    )


# ── Images ──

def _extract_image(
    file_path: Path, use_ocr: bool = True, ocr_engine: str = "tesseract",
    force_ocr: bool = False,
) -> ExtractedDocument:
    """OCR an image file directly."""
    from PIL import Image
    from core.ocr import ocr_page

    warnings = []
    img = Image.open(str(file_path))

    if not use_ocr:
        return ExtractedDocument(
            source_file=file_path.name,
            raw_text="",
            page_count=1,
            extraction_method="none",
            warnings=["OCR disabled — no text extracted from image"],
        )

    try:
        text = ocr_page(img, engine=ocr_engine)
    except Exception as e:
        text = ""
        warnings.append(f"OCR failed: {e}")

    return ExtractedDocument(
        source_file=file_path.name,
        raw_text=text,
        page_count=1,
        extraction_method=f"ocr_{ocr_engine}",
        warnings=warnings,
    )


# ── Dispatcher ──

_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".doc": _extract_docx,
    ".xlsx": _extract_xlsx,
    ".xls": _extract_xlsx,
    ".csv": _extract_csv,
    ".txt": _extract_txt,
    ".md": _extract_txt,
    ".png": _extract_image,
    ".jpg": _extract_image,
    ".jpeg": _extract_image,
    ".tiff": _extract_image,
    ".tif": _extract_image,
    ".bmp": _extract_image,
    ".webp": _extract_image,
}
